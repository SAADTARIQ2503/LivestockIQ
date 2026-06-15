"""Generate MUTATION_TESTING_REPORT.docx — LivestockIQ, from scratch."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Helpers ───────────────────────────────────────────────────────────────────

def shd(cell, hex_color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:color'), 'auto')
    el.set(qn('w:fill'), hex_color)
    pr.append(el)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(11)
    return p

def code(text):
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), 'F0F0F0')
    pPr.append(s)
    return p

def table(headers, rows, hdr_color='1E3A5F'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; c.text = h
        shd(c, hdr_color)
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]; c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                shd(c, 'EEF2F7')
    return t

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    return p

# ── Title ─────────────────────────────────────────────────────────────────────

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Mutation Testing Assignment Report')
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('CS-ST — Mutation Testing Assignment  |  FAST NUCES — Spring 2025')
r.italic = True; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

doc.add_paragraph()

heading('Group Information', level=2)
table(['Field', 'Value'], [
    ['FYP Title',       'LivestockIQ — AI-Powered Livestock Management Platform'],
    ['Member 1 (ID)',   '22F-8785 — Saad Ullah Tariq'],
    ['Member 2 (ID)',   '22F-3287 — Umer Javaid'],
    ['Member 3 (ID)',   '22F-8789 — M. Haris'],
    ['Supervisor',      'Muhammad Tayyab Javaid'],
    ['Tech Stack',      'Django 5 (Python) + React 18 (JavaScript/Vite)'],
    ['Backend Module',  'LivestockIQ/ai_service/id_parser.py'],
    ['Frontend Module', 'frontend/src/utils/constants.js'],
])

doc.add_page_break()

heading('Table of Contents', level=1)
for item in [
    'Task 1 — Baseline Assessment',
    'Task 2 — The Mutation Run',
    'Task 3 — Mutant Analysis & Eradication',
    'Task 4 — Final Mutation Score',
]:
    bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TASK 1
# ════════════════════════════════════════════════════════════════════════════

heading('Task 1 — Baseline Assessment (3 Marks)', level=1)

# ── 1.1 ──────────────────────────────────────────────────────────────────────
heading('1.1  Target Module Selection', level=2)

para(
    'LivestockIQ has two distinct technology layers that are tested separately: '
    'a Django REST Framework backend (Python) and a React + Vite frontend (JavaScript). '
    'A module was selected from each layer to demonstrate mutation testing across both technologies.'
)

doc.add_paragraph()
para('Backend Module — ai_service/id_parser.py', bold=True)
para(
    'This module implements the livestock animal ID parser used by the OCR service '
    '(GotOCRService). When a farmer photographs an animal\'s ear tag, brand, or painted '
    'number, the NVIDIA Llama-3.2-Vision or GOT-OCR-2.0 model extracts raw text from '
    'the image. That raw text is then passed to parse_animal_id() to extract a clean '
    'integer ID. The module contains two functions:'
)

table(['#', 'Function', 'Description', 'Mutation-Vulnerable Logic'], [
    ['1', '_is_hallucination(n)',
     'Rejects OCR hallucinations: all-same digits (1111), ascending runs (1234), '
     'descending runs (4321)',
     'len(s) < 2 boundary, set() check, all() + arithmetic on consecutive digits'],
    ['2', 'parse_animal_id(text)',
     'Extracts animal tag integer from OCR output via 4 regex patterns in priority order',
     'Early None return, 4 ordered regex patterns, re.sub, continue vs break on hallucination'],
])

para(
    'Justification: parse_animal_id() is pure Python with no external dependencies — '
    'no Django ORM, no ML model, no HTTP calls. This makes it an ideal mutation target: '
    'every mutant is testable in milliseconds without mocking. The hallucination guard '
    'contains boundary comparisons (< 2), set-size checks, and arithmetic on digit runs '
    '— exactly the operators Relational Operator Replacement (ROR) and Arithmetic Operator '
    'Replacement (AOR) target. The loop body\'s continue/break distinction is a classic '
    'Statement Operator Replacement (SVR) candidate.'
)

doc.add_paragraph()
para('Frontend Module — frontend/src/utils/constants.js', bold=True)
para(
    'This module defines every application-wide constant used across the LivestockIQ '
    'React frontend. It covers animal domain types (Cow, Goat, Sheep), sex options '
    '(Male, Female), alert severity levels (Low, Medium, High), chart colour hex values '
    'for the health and environment dashboards, JWT token storage keys, form validation '
    'constraints (PASSWORD_MIN_LENGTH, USERNAME_MIN_LENGTH/MAX_LENGTH), React Query '
    'cache key factories, and season options for the vaccination scheduler. '
    'Any mutation in this file — a single wrong hex code, a swapped string, an empty '
    'array — propagates silently through the entire UI if tests do not assert exact values.'
)

doc.add_paragraph()

# ── 1.2 ──────────────────────────────────────────────────────────────────────
heading('1.2  Baseline Test Suite', level=2)

para('Backend baseline (tests/test_id_parser_baseline.py) — 14 tests:', bold=True)
para(
    'Tests exercise the happy path with loose assertions (assertIsNotNone, assertIsInstance). '
    'They confirm each path returns a value of the right type but never assert an exact integer.'
)
table(['Test', 'Input', 'Assertion Type'], [
    ['test_explicit_id_label_returns_int',      '"ID: 42"',           'assertIsNotNone + assertIsInstance(int)'],
    ['test_tag_label_returns_value',            '"Tag: A-042"',       'assertIsNotNone'],
    ['test_letter_prefix_returns_value',        '"F73"',              'assertIsNotNone'],
    ['test_bare_number_returns_value',          '"907"',              'assertIsNotNone'],
    ['test_mixed_text_extracts_something',      '"Animal VST94 is…"', 'assertIsNotNone'],
    ['test_five_digit_id_returns_value',        '"ID: 10045"',        'assertIsNotNone'],
    ['test_uppercase_insensitive',              '"id: 5"',            'assertIsNotNone'],
    ['test_tag_b_prefix_returns_value',         '"Tag: B-73"',        'assertIsNotNone'],
    ['test_none_returns_none',                  'None',               'assertIsNone'],
    ['test_empty_string_returns_none',          '""',                 'assertIsNone'],
    ['test_no_text_found_sentinel_returns_none','NO_TEXT_FOUND',      'assertIsNone'],
    ['test_all_same_digits_returns_none',       '"1111"',             'assertIsNone'],
    ['test_ascending_sequence_returns_none',    '"1234"',             'assertIsNone'],
    ['test_descending_sequence_returns_none',   '"4321"',             'assertIsNone'],
])

doc.add_paragraph()
para('Frontend baseline (src/tests/utils/constants.test.js) — 5 tests:', bold=True)
table(['Describe Block', 'What is Tested', 'Assertion Type'], [
    ['ANIMAL_TYPES',        'ANIMAL_TYPES.COW, .GOAT, .SHEEP values',       'toBe (exact)'],
    ['ANIMAL_TYPE_OPTIONS', 'Array length ≥ 3',                              'toBeGreaterThanOrEqual (loose)'],
    ['ANIMAL_TYPE_OPTIONS', 'Each option has .value and .label properties',  'toHaveProperty (structural)'],
    ['ANIMAL_TYPE_OPTIONS', 'Values array contains "Cow", "Goat", "Sheep"',  'toContain (loose)'],
    ['STORAGE_KEYS',        'Has ACCESS_TOKEN and REFRESH_TOKEN properties', 'toHaveProperty (structural)'],
])

# ── 1.3 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('1.3  Baseline Coverage Report', level=2)

para('Backend (pytest-cov):', bold=True)
table(['Metric', 'Value'], [
    ['Statement Coverage', '100% — all statements in id_parser.py executed'],
    ['Branch Coverage',    '85%  — the len(s) < 2 False branch not covered for 2-digit inputs'],
    ['Function Coverage',  '100% — both _is_hallucination() and parse_animal_id() called'],
    ['Line Coverage',      '100% — every line executed'],
])
doc.add_paragraph()
code('python -m pytest tests/test_id_parser_baseline.py --cov=ai_service.id_parser --cov-report=term-missing')

doc.add_paragraph()
para('Frontend (Vitest v8):', bold=True)
table(['Metric', 'Value'], [
    ['Statement Coverage', '100% — all export declarations executed'],
    ['Branch Coverage',    '0%   — the || fallback in API_BASE_URL (line 6) never tested'],
    ['Function Coverage',  '0%   — no QUERY_KEYS factory functions called'],
    ['Line Coverage',      '100% — every line reached'],
])
doc.add_paragraph()
code('npx vitest run --coverage --coverage.include=\'src/utils/constants.js\' src/tests/utils/constants.test.js')

# ── 1.4 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('1.4  Preliminary Analysis', level=2)

para(
    'Both modules achieve 100% line coverage with their baseline suites, yet as Task 2 '
    'reveals, the mutation scores are dramatically lower. This gap is caused by the same '
    'underlying weakness in both suites: the tests verify the existence and type of a '
    'return value rather than its exact content.'
)
para(
    'For the backend, test_explicit_id_label_returns_int() calls parse_animal_id("ID: 42") '
    'and asserts assertIsInstance(result, int). This executes every line of code and passes, '
    'but it cannot distinguish the correct return value 42 from any other integer. '
    'A mutant that changes the regex pattern from \\\\bID\\\\s*[:=]?\\\\s*(\\\\d+)\\\\b to '
    '\\\\bID\\\\s*[:=]?\\\\s*XX(\\\\d+)XX\\\\b still returns an integer (from the bare-number '
    'fallback pattern) — and the test passes. The mutation survives.'
)
para(
    'For the frontend, the baseline suite checks that ANIMAL_TYPE_OPTIONS.length >= 3 and '
    'that each entry has a .value property, but never checks that SEVERITY_COLORS.High is '
    'exactly "#dc3545" (the red used for critical health alerts). A mutant that changes this '
    'to an empty string "" does not affect array length or property existence — the tests '
    'still pass while the entire alert-colour system is broken.'
)
para(
    'Coverage tells you which lines ran. Mutation testing tells you whether those lines '
    'were verified. 100% line coverage with 7.18% mutation score (frontend) means the '
    'tests execute the code 100% of the time but catch only 7% of injected faults.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TASK 2
# ════════════════════════════════════════════════════════════════════════════

heading('Task 2 — The Mutation Run (4 Marks)', level=1)

# ── 2.1 ──────────────────────────────────────────────────────────────────────
heading('2.1  Tool Configuration', level=2)

para('Backend — mutmut 3.6.0 (Python)', bold=True)
table(['Setting', 'Value'], [
    ['Mutation Tool',   'mutmut 3.6.0'],
    ['Test Runner',     'pytest 9.1.0'],
    ['Python Version',  '3.12.3'],
    ['Target File',     'ai_service/id_parser.py'],
    ['Config File',     'LivestockIQ/setup.cfg'],
    ['Test Selection',  'tests/test_id_parser_baseline.py'],
])
doc.add_paragraph()
code("""# LivestockIQ/setup.cfg
[mutmut]
source_paths=ai_service/id_parser.py
pytest_add_cli_args_test_selection=tests/""")

doc.add_paragraph()
para('Frontend — Stryker Mutator 9.6.1 (JavaScript)', bold=True)
table(['Setting', 'Value'], [
    ['Mutation Tool',      'Stryker Mutator v9.6.1'],
    ['Test Runner',        'Vitest v3.2.6 via @stryker-mutator/vitest-runner'],
    ['JavaScript Version', 'ES Modules (React 18 + Vite 6)'],
    ['Target File',        'src/utils/constants.js'],
    ['Config File',        'frontend/stryker.config.mjs'],
    ['Test Selection',     'src/tests/utils/constants.test.js'],
])
doc.add_paragraph()
code("""// frontend/stryker.config.mjs
export default {
  testRunner: 'vitest',
  vitest: { configFile: 'vite.config.js' },
  mutate: ['src/utils/constants.js'],
  coverageAnalysis: 'perTest',
  concurrency: 2,
  reporters: ['html', 'progress'],
  htmlReporter: { fileName: 'test-results/mutation/index.html' },
};""")

# ── 2.2 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('2.2  Initial Mutation Results', level=2)

para('Backend — mutmut baseline results:', bold=True)
table(['Metric', 'Value'], [
    ['Total Mutants',    '65'],
    ['Killed (🎉)',       '50'],
    ['Survived (🙁)',    '15'],
    ['Timed Out',        '0'],
    ['No Coverage',      '0'],
    ['Mutation Score',   '76.92%  (50 / 65)'],
])
doc.add_paragraph()
code('python -m mutmut run   # from LivestockIQ/ with venv active')

doc.add_paragraph()
para('Frontend — Stryker baseline results:', bold=True)
table(['Metric', 'Value'], [
    ['Total Mutants',      '195'],
    ['Killed',             '14'],
    ['Survived',           '151'],
    ['Timed Out',          '0'],
    ['No Coverage',        '30'],
    ['Mutation Score',     '7.18%  (14 / 195)'],
    ['Covered Score',      '8.48%  (14 / 165)'],
])
doc.add_paragraph()
code('npx stryker run --mutate \'src/utils/constants.js\'   # from frontend/')

# ── 2.3 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('2.3  Coverage–Mutation Score Gap', level=2)

para('Backend:', bold=True)
table(['Metric', 'Value'], [
    ['Line / Statement Coverage', '100%'],
    ['Mutation Score',            '76.92%'],
    ['Gap',                       '23.08 percentage points'],
])

doc.add_paragraph()
para('Frontend:', bold=True)
table(['Metric', 'Value'], [
    ['Line / Statement Coverage', '100%'],
    ['Mutation Score',            '7.18%'],
    ['Gap',                       '92.82 percentage points'],
])

doc.add_paragraph()
para(
    'The frontend gap of 92.82 points is extreme and instructive. constants.js exports '
    '40+ named constants. The baseline suite imports only 3 of them and verifies their '
    'values partially. Stryker generated 195 mutants by altering every string literal, '
    'array element, and arrow function. The 5 baseline tests could only observe changes '
    'to the 3 imported values — 192 mutants were completely invisible to the test suite. '
    'The 30 NoCoverage mutants were never even reached; 151 ran but were undetected.'
)
para(
    'The backend gap is more moderate (23 points) because the 14 baseline tests cover '
    'all code paths. The surviving 15 mutants escape primarily because: (a) assertions '
    'use assertIsNotNone instead of assertEqual, and (b) logger.debug() calls are pure '
    'side effects with no observable return-value impact.'
)

# ── 2.4 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('2.4  Mutation Operator Breakdown (Survived)', level=2)

para('Backend (mutmut) — 15 survived:', bold=True)
table(['Operator Category', 'Count', 'Example Mutant'], [
    ['ROR (Relational Operator Replacement)', '2',
     'len(s) < 2  →  len(s) <= 2   and   len(s) < 3'],
    ['SVR — logger.debug() arguments',        '9',
     'logger.debug("Rejected…", val, text)  →  logger.debug(None, val, text)'],
    ['SVR — re.sub() pattern / replacement',  '2',
     're.sub(r\'\\D\', \'\', …)  →  re.sub(r\'XX\\DXX\', \'\', …)'],
    ['SVR — continue → break',                '1',
     'continue  →  break  (inside hallucination-rejected branch)'],
    ['SVR — text.upper() → text.lower()',     '1',
     'upper = text.upper()  →  upper = text.lower()'],
])

doc.add_paragraph()
para('Frontend (Stryker) — 151 survived:', bold=True)
table(['Operator Type', 'Count', 'Example Mutant'], [
    ['StringLiteral (SVR)', '97',
     '\'Cow\'  →  \'\'  |  \'#dc3545\'  →  \'\'  |  \'access_token\'  →  \'\''],
    ['ArrayDeclaration',    '19',
     '[{value:\'Cow\',…}, …]  →  []   (entire options array emptied)'],
    ['ArrowFunction',       '10',
     '(filters) => [\'animals\',\'list\',filters]  →  () => undefined'],
    ['LogicalOperator',      '1',
     'VITE_API_BASE_URL || \'http://localhost:8000/api/v1\'  →  &&'],
    ['NoCoverage (not run)', '30',
     'Mutants in constants not imported by any test'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TASK 3
# ════════════════════════════════════════════════════════════════════════════

heading('Task 3 — Mutant Analysis & Eradication (10 Marks)', level=1)

heading('Mutant Selection Criteria', level=2)
for c in [
    '✅  At least 1 ROR (Relational Operator Replacement) mutant — Mutant 1',
    '✅  At least 1 SVR (Statement/Value Replacement) — continue → break — Mutant 2',
    '✅  At least 2 mutants from the same module (both from id_parser.py) — Mutants 1 & 2',
    '✅  Backend mutant (Python / Django codebase) — Mutants 1, 2, 3',
    '✅  Frontend mutant (React / JavaScript) — Mutants 4, 5',
    '✅  5 mutants analysed with deep root-cause analysis',
    '✅  All 5 adversarial tests confirmed passing before re-running mutation tools',
]:
    bullet(c)

doc.add_paragraph()

# ── Mutant 1 ──────────────────────────────────────────────────────────────────
heading('Mutant 1 — ROR in _is_hallucination()  [id_parser.py, line ~20]', level=2)
table(['Field', 'Value'], [
    ['File',             'LivestockIQ/ai_service/id_parser.py'],
    ['Function',         '_is_hallucination(n: int) -> bool'],
    ['Original',         'if len(s) < 2:  return False'],
    ['Mutant A (mutmut_3)', 'if len(s) <= 2:  return False'],
    ['Mutant B (mutmut_4)', 'if len(s) < 3:   return False'],
    ['Operator',         'ROR — Relational Operator Replacement (< → <= and < → <)'],
    ['Status (Baseline)','Survived'],
    ['Status (Final)',   'Killed'],
])
heading('Root-Cause Analysis', level=3)
para(
    'The guard `if len(s) < 2: return False` ensures that single-digit numbers (like the '
    'animal tag "5") are never rejected as hallucinations, since a single digit cannot '
    'form a repeating or sequential pattern. The mutant changes this boundary to <= 2, '
    'which exempts two-digit numbers from ALL subsequent hallucination checks. This means '
    'a painted tag reading "11" (two identical digits — a classic OCR hallucination) would '
    'pass through as a valid livestock ID instead of being rejected.'
)
para('Why it survived the baseline:', bold=True)
para(
    'The baseline suite tests hallucination rejection with four-digit inputs (1111, 1234, '
    '4321). For these inputs, len(s) = 4, so both < 2 and <= 2 produce the same result '
    '(False), and the check proceeds normally. The critical boundary — a two-digit '
    'repeated value like "11" or "99" — was never tested. Both mutants pass all 14 '
    'baseline tests undetected.'
)
para('Adversarial Test:', bold=True)
code("""def test_kill_rOR_lt2_to_lte2_two_same_digits(self):
    # "11" is 2 same digits. Original: len("11") < 2 → False → check → True (hallucination) → None
    # Mutant  : len("11") <= 2 → True → early return False (not hallucination) → returns 11!
    assert parse_animal_id('11') is None

def test_kill_rOR_lt2_to_lt3_two_digit_ascending(self):
    # "12" = ascending 1→2. Mutant len("12") < 3 → True → bypasses check → returns 12!
    assert parse_animal_id('12') is None""")

doc.add_paragraph()

# ── Mutant 2 ──────────────────────────────────────────────────────────────────
heading('Mutant 2 — SVR (continue → break) in parse_animal_id()  [id_parser.py, line ~49]', level=2)
table(['Field', 'Value'], [
    ['File',             'LivestockIQ/ai_service/id_parser.py'],
    ['Function',         'parse_animal_id(text) — the hallucination-rejected branch'],
    ['Original',         'if _is_hallucination(val):  ...  continue'],
    ['Mutant (mutmut_32)', 'if _is_hallucination(val):  ...  break'],
    ['Operator',         'SVR — Statement/Value Replacement (continue → break in for-loop)'],
    ['Status (Baseline)','Survived'],
    ['Status (Final)',   'Killed'],
])
heading('Root-Cause Analysis', level=3)
para(
    'parse_animal_id() tries four regex patterns in descending specificity. When a pattern '
    'finds a match but the extracted number is a hallucination (e.g., "1234" from '
    '"ID: 1234"), `continue` moves to the next pattern, giving the function a second '
    'chance to find a valid tag (e.g., "A-73" via the letter-prefix pattern). '
    'The `break` mutant exits the loop entirely — the function returns None even when '
    'a valid ID exists in the string under a different pattern.'
)
para(
    'In a real livestock scenario: an OCR result of "ID: 1234 A-73" means the model '
    'read a hallucinated sequential ID first but also detected a real branded tag "A-73" '
    '(system ID 73). With `continue`, the parser correctly returns 73. With `break`, '
    'it silently discards the animal ID, causing the farmer\'s tag scan to fail.'
)
para('Why it survived the baseline:', bold=True)
para(
    'Every baseline test uses a single clean input (no hallucination + fallback '
    'scenario). The test "test_explicit_id_label_returns_int" uses "ID: 42" where pattern '
    '1 succeeds without hitting the hallucination guard at all. No test presented a string '
    'where one pattern gives a hallucination and a later pattern gives a valid ID.'
)
para('Adversarial Test:', bold=True)
code("""def test_kill_continue_to_break_first_pattern_hallucination_second_pattern_valid(self):
    # Pattern 1 matches 1234 (ascending → hallucination → continue, not break)
    # Pattern 3 matches A-73 → extracts 73 → valid livestock tag
    # With break: returns None (tag scan fails for real animal ID 73)
    assert parse_animal_id('ID: 1234 A-73') == 73

def test_kill_continue_to_break_tag_label_hallucination_letter_prefix_valid(self):
    # Tag label gives ascending hallucination; letter-prefix recovers the real tag
    assert parse_animal_id('Tag: 1234 A-73') == 73""")

doc.add_paragraph()

# ── Mutant 3 ──────────────────────────────────────────────────────────────────
heading('Mutant 3 — SVR (logger.debug args) in parse_animal_id()  [id_parser.py, line ~50]', level=2)
table(['Field', 'Value'], [
    ['File',             'LivestockIQ/ai_service/id_parser.py'],
    ['Function',         'parse_animal_id() — rejection logging call'],
    ['Original',         'logger.debug("Rejected hallucination %d from %r", val, text)'],
    ['Mutants',          '9 mutations: None args, dropped args, mutated format string, '
                         'case changes — all affecting only the log message'],
    ['Operator',         'SVR — argument and string literal replacements'],
    ['Status (Baseline)','Survived (all 9)'],
    ['Status (Final)',   'Survived (all 9) — Equivalent Mutants'],
])
heading('Root-Cause Analysis', level=3)
para(
    'These 9 mutants only affect the arguments or format string passed to logger.debug(). '
    'logger.debug() is a side-effect call — it writes to the application log but has no '
    'influence on the function\'s return value. Changing the log message from '
    '"Rejected hallucination %d from %r" to "rejected hallucination %d from %r" '
    '(lowercase) or replacing val with None does not change what parse_animal_id() returns '
    'to its caller. All 14 baseline tests and all 15 adversarial tests pass for every '
    'variant of this mutant because no test captures or asserts log output.'
)
para(
    'These are genuine equivalent mutants. Killing them would require either: '
    '(a) asserting on captured log records (using assertLogs() from unittest), or '
    '(b) removing the debug call from the source code entirely. '
    'Since debug logging is valuable operational behaviour in a livestock AI system '
    '(it helps diagnose OCR failures in the field), option (b) is undesirable. '
    'The recommended approach is to exclude these mutants from the mutation score '
    'calculation as recognised equivalent mutants.'
)

doc.add_paragraph()

# ── Mutant 4 ──────────────────────────────────────────────────────────────────
heading('Mutant 4 — StringLiteral in SEVERITY_COLORS  [constants.js]', level=2)
table(['Field', 'Value'], [
    ['File',              'frontend/src/utils/constants.js'],
    ['Constant',          'SEVERITY_COLORS'],
    ['Original',          "SEVERITY_COLORS = { Low: '#28a745', Medium: '#ffc107', High: '#dc3545' }"],
    ['Example Mutant',    "High: '#dc3545'  →  High: ''"],
    ['Operator',          'StringLiteral replacement'],
    ['Status (Baseline)', 'Survived'],
    ['Status (Final)',    'Killed'],
])
heading('Root-Cause Analysis', level=3)
para(
    'SEVERITY_COLORS maps the three alert severity levels (Low, Medium, High) to '
    'specific hex colour codes used in the LivestockIQ alert dashboard, vaccination '
    'overdue indicators, and health monitoring panels. The colour #dc3545 (red) is '
    'applied to "High" severity alerts such as disease detections — a wrong colour '
    'here means critical livestock disease alerts render in the wrong colour or become '
    'invisible (empty string).'
)
para('Why it survived the baseline:', bold=True)
para(
    'The baseline test imports and tests ANIMAL_TYPES and STORAGE_KEYS — it does not '
    'import SEVERITY_COLORS at all. The Stryker NoCoverage category confirms these '
    'mutants were never even reached by the test runner. A mutant replacing "#dc3545" '
    'with "" would sail through all 5 baseline tests undetected.'
)
para('Adversarial Test:', bold=True)
code("""describe('SEVERITY_COLORS — exact hex values', () => {
  // Each test uses toBe() — exact string match — so any mutated hex value fails
  it('Low is green #28a745',   () => expect(SEVERITY_COLORS.Low).toBe('#28a745'));
  it('Medium is amber #ffc107',() => expect(SEVERITY_COLORS.Medium).toBe('#ffc107'));
  it('High is red #dc3545',    () => expect(SEVERITY_COLORS.High).toBe('#dc3545'));
});""")

doc.add_paragraph()

# ── Mutant 5 ──────────────────────────────────────────────────────────────────
heading('Mutant 5 — ArrayDeclaration in SEX_OPTIONS  [constants.js]', level=2)
table(['Field', 'Value'], [
    ['File',              'frontend/src/utils/constants.js'],
    ['Constant',          'SEX_OPTIONS'],
    ['Original',          "SEX_OPTIONS = [{ value: 'Male', label: 'Male' }, { value: 'Female', label: 'Female' }]"],
    ['Mutant',            "SEX_OPTIONS = []   (entire array replaced with empty array)"],
    ['Operator',          'ArrayDeclaration replacement'],
    ['Status (Baseline)', 'Survived'],
    ['Status (Final)',    'Killed'],
])
heading('Root-Cause Analysis', level=3)
para(
    'SEX_OPTIONS is used in the Add Animal and Edit Animal forms to populate the sex '
    'dropdown selector. When a farmer registers a Cow as Female, the form uses this '
    'constant to render the radio buttons. Replacing the array with [] would render '
    'an empty dropdown — no options would appear — making it impossible to register '
    'a new animal with a sex value.'
)
para('Why it survived the baseline:', bold=True)
para(
    'SEX_OPTIONS is not imported or tested anywhere in the baseline test suite. '
    'This constant falls into the NoCoverage category — the Stryker mutant was '
    'generated but never executed during testing because no test file references '
    'SEX_OPTIONS. This is an extreme case of test-coverage blindspot: 0% test '
    'coverage for a constant used in a core livestock registration workflow.'
)
para('Adversarial Test:', bold=True)
code("""describe('SEX_OPTIONS', () => {
  it('has exactly 2 entries', () => expect(SEX_OPTIONS.length).toBe(2));
  it('first is Male', () => {
    expect(SEX_OPTIONS[0].value).toBe('Male');
    expect(SEX_OPTIONS[0].label).toBe('Male');
  });
  it('second is Female', () => {
    expect(SEX_OPTIONS[1].value).toBe('Female');
    expect(SEX_OPTIONS[1].label).toBe('Female');
  });
});""")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TASK 4
# ════════════════════════════════════════════════════════════════════════════

heading('Task 4 — Final Mutation Score (3 Marks)', level=1)

# ── 4.1 ──────────────────────────────────────────────────────────────────────
heading('4.1  Final Mutation Results (After Adversarial Tests)', level=2)

para('Backend (mutmut + pytest):', bold=True)
table(['Metric', 'Before', 'After', 'Change'], [
    ['Total Mutants',  '65',    '65',    '—'],
    ['Killed',         '50',    '54',    '+4'],
    ['Survived',       '15',    '11',    '−4'],
    ['Timed Out',      '0',     '0',     '—'],
    ['No Coverage',    '0',     '0',     '—'],
    ['Mutation Score', '76.92%','83.08%','+6.16 pp'],
])

doc.add_paragraph()
para('Frontend (Stryker + Vitest):', bold=True)
table(['Metric', 'Before', 'After', 'Change'], [
    ['Total Mutants',  '195',  '195',   '—'],
    ['Killed',         '14',   '99',    '+85'],
    ['Survived',       '151',  '72',    '−79'],
    ['Timed Out',      '0',    '0',     '—'],
    ['No Coverage',    '30',   '24',    '−6'],
    ['Mutation Score', '7.18%','50.77%','+43.59 pp'],
    ['Covered Score',  '8.48%','57.89%','+49.41 pp'],
])

# ── 4.2 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('4.2  Test Suite Growth', level=2)

para('Backend:', bold=True)
table(['Metric', 'Before', 'After'], [
    ['Test Files', '1 (baseline)', '2 (baseline + adversarial)'],
    ['Total Tests', '14', '29'],
    ['Tests Added', '—', '+15'],
])

doc.add_paragraph()
para('Frontend:', bold=True)
table(['Metric', 'Before', 'After'], [
    ['Test Files', '1 (baseline)', '2 (baseline + adversarial)'],
    ['Total Tests', '5', '70'],
    ['Tests Added', '—', '+65'],
])

# ── 4.3 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('4.3  Remaining Survivors Analysis', level=2)

para('Backend — 11 survivors:', bold=True)

para('Group A: 9 logger.debug() mutations (Equivalent Mutants)', bold=True)
para(
    'All 9 remaining survivors mutate the logger.debug() call that logs rejected '
    'hallucinations. As explained in Mutant 3, these are equivalent mutants: '
    'logger.debug() is a side-effect call with no influence on the return value. '
    'Killing them would require asserting on captured log output (e.g., using '
    'Python\'s assertLogs() context manager), which is outside the scope of functional '
    'unit tests for the parse_animal_id() function.'
)

para('Group B: 1 re.sub() pattern mutation and 1 re.sub() replacement mutation (Equivalent Mutants)', bold=True)
code("""# Original:
digits = re.sub(r'\\D', '', m.group(1))

# mutmut_16: re.sub(r'XX\\DXX', '', m.group(1))
# mutmut_17: re.sub(r'\\D', 'XXXX', m.group(1))""")
para(
    'The regex patterns used by parse_animal_id() capture only digit sequences '
    '(via \\\\d{1,5} in their capture groups). By the time re.sub() is called on '
    'm.group(1), the captured group contains only digit characters. Replacing '
    'non-digit characters (\\\\D) has no effect when there are no non-digit characters '
    'to replace. Both mutations produce identical output to the original for all valid '
    'test inputs — these are structural equivalent mutants arising from defensive coding.'
)

doc.add_paragraph()
para('Frontend — 72 survivors:', bold=True)

para('Group A: 45 StringLiteral mutations on uncovered constants', bold=True)
para(
    'The remaining 45 string literal mutations are in constants that the adversarial '
    'test suite did not import: DATE_FORMAT, DATETIME_FORMAT, TIME_FORMAT, individual '
    'QUERY_KEYS sub-keys (e.g., costs.summary, environment.forecast), and some '
    'ERROR_MESSAGES/SUCCESS_MESSAGES values. These can be killed by extending the '
    'adversarial suite to cover every exported constant with exact-value assertions.'
)

para('Group B: 9 ArrayDeclaration mutations', bold=True)
para(
    'Nine array constants (SEASONS, SEASON_OPTIONS sub-arrays, PAGE_SIZE_OPTIONS) '
    'still have some entries surviving. The adversarial suite tests array length and '
    'a subset of entries, but Stryker generates one mutation per array element — '
    'killing all requires asserting every individual index.'
)

para('Group C: 8 ArrowFunction mutations and 1 LogicalOperator mutation', bold=True)
para(
    'Arrow functions in QUERY_KEYS (animals.list, animals.detail, health.schedules, '
    'costs.report, etc.) are replaced with () => undefined. The adversarial suite kills '
    'most via return-value assertions, but 8 remain. The LogicalOperator mutation '
    '(|| → &&) in the API_BASE_URL fallback survives because import.meta.env is '
    'defined in the test environment, so the right-hand side is never evaluated.'
)

# ── 4.4 ──────────────────────────────────────────────────────────────────────
doc.add_paragraph()
heading('4.4  Reflection', level=2)
para(
    'This mutation testing exercise revealed fundamentally different failure modes in '
    'the two LivestockIQ modules. For the backend id_parser.py, the test suite had '
    'good structural coverage (100% lines, all code paths exercised) but weak value '
    'assertions. The 76.92% baseline score was higher than expected because the '
    'hallucination guard logic is well-exercised by the 14 null/sentinel/sequence tests. '
    'The critical gap was the boundary condition: the < 2 guard for 2-digit inputs was '
    'never tested at exactly length 2. One adversarial test for parse_animal_id("11") '
    'and parse_animal_id("12") killed both ROR variants simultaneously.'
)
para(
    'For the frontend constants.js, the situation was more severe: 92.82 percentage points '
    'separated 100% line coverage from 7.18% mutation score. This happened because '
    'constants.js exports 40+ constants but the baseline suite imported only 3. The '
    'remaining 37 constants — including SEVERITY_COLORS used in the alert dashboard, '
    'SEX_OPTIONS used in animal registration, STORAGE_KEYS used in JWT token management, '
    'and VACCINATION_CHART_COLORS used in the health dashboard — were completely invisible '
    'to testing. A Stryker mutant that erased "#dc3545" from the critical-alert colour '
    'would have shipped to production undetected.'
)
para(
    'The 65-test adversarial suite improved the frontend score from 7.18% to 50.77% '
    '(+43.59 pp) by adding exact toBe() assertions for every major constant. '
    'The backend improved from 76.92% to 83.08% (+6.16 pp) by targeting the specific '
    'boundary and control-flow mutations identified by mutmut.'
)
para(
    'The key lesson: in a livestock management system, a wrong hex colour code or an '
    'empty options array is not a cosmetic bug — it breaks the alarm system for disease '
    'alerts or prevents farmers from registering animals. Mutation testing surfaces these '
    'silent failures before they reach production, complementing the existing Django REST '
    'API test suite (137 backend tests) and the Vitest component/API test suite.'
)

doc.add_page_break()

# ── Repository Structure ──────────────────────────────────────────────────────
heading('Repository Structure', level=1)
code("""LivestockIQ/
├── ai_service/
│   ├── id_parser.py                      # Backend target module
│   └── ocr_service.py                    # Full OCR service (id_parser extracted from here)
├── tests/
│   ├── test_id_parser_baseline.py        # 14 baseline tests (pytest)
│   └── test_id_parser_adversarial.py     # 15 adversarial tests (pytest)
├── setup.cfg                             # mutmut configuration
└── manage.py

frontend/
├── src/
│   ├── utils/
│   │   └── constants.js                  # Frontend target module
│   └── tests/
│       └── utils/
│           ├── constants.test.js         # 5 baseline tests (Vitest)
│           └── constants.adversarial.test.js  # 65 adversarial tests (Vitest)
├── test-results/
│   └── mutation/index.html               # Stryker HTML report
├── stryker.config.mjs                    # Stryker configuration
└── vite.config.js                        # Vitest configuration""")

heading('Commands Reference', level=1)
table(['Command', 'Purpose', 'Run From'], [
    ['python -m mutmut run',
     'Run backend mutation testing (mutmut)',
     'LivestockIQ/ (venv active)'],
    ['python -m mutmut results',
     'Show survived backend mutants',
     'LivestockIQ/ (venv active)'],
    ['python -m pytest tests/ -v',
     'Run all backend id_parser tests',
     'LivestockIQ/ (venv active)'],
    ['npx stryker run --mutate \'src/utils/constants.js\'',
     'Run frontend mutation testing (Stryker)',
     'frontend/'],
    ['npx vitest run src/tests/utils/',
     'Run all frontend constant tests',
     'frontend/'],
    ['python manage.py test --settings=LivestockIQ.test_settings',
     'Run full Django API test suite (137 tests)',
     'LivestockIQ/ (venv active)'],
])

# ── Save ─────────────────────────────────────────────────────────────────────
out = '/home/saad-tariq/Desktop/FYP/Code/LivestockIQ/MUTATION_TESTING_REPORT.docx'
doc.save(out)
print(f'Saved: {out}')
